import os
import uuid
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
import docx.shared
import re
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, jsonify, make_response
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from flask_cors import CORS
import pytesseract
from PIL import Image
import io
import zipfile
import json

# Configure Tesseract path - adjust this to your actual installation path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf'}
app.config['ALLOWED_IMAGE_EXTENSIONS'] = {'jpg', 'jpeg', 'png'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

# Create uploads directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure Google Generative AI
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    print("WARNING: GOOGLE_API_KEY environment variable not set. Gemini features will not work.")
else:
    genai.configure(api_key=GOOGLE_API_KEY)

def allowed_file(filename):
    """Check if the uploaded file has an allowed extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def allowed_image_file(filename):
    """Check if the uploaded file is an allowed image extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_IMAGE_EXTENSIONS']

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using PyMuPDF, with OCR fallback for scanned documents, and save images"""
    text_content = []
    image_paths = []  # List to store paths of extracted images
    try:
        # Create a directory for images if it doesn't exist
        base_dir = os.path.dirname(pdf_path)
        images_dir = os.path.join(base_dir, "extracted_images")
        os.makedirs(images_dir, exist_ok=True)
        
        # Generate a unique prefix for this PDF's images
        image_prefix = os.path.splitext(os.path.basename(pdf_path))[0] + "_"
        
        # Open the PDF
        doc = fitz.open(pdf_path)
        
        # Extract text and images from each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            
            # Extract images from the page
            image_list = page.get_images(full=True)
            
            # Process each image
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]  # Image reference
                
                # Try to get the image
                try:
                    base_img = doc.extract_image(xref)
                    image_bytes = base_img["image"]
                    image_ext = base_img["ext"]
                    
                    # Save the image to a file
                    img_filename = f"{image_prefix}page{page_num+1}_img{img_index+1}.{image_ext}"
                    img_path = os.path.join(images_dir, img_filename)
                    
                    with open(img_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    # Add image path to our list
                    image_paths.append(img_path)
                    
                    # Add a placeholder in the text for this image
                    text_content.append(f"[IMAGE: {img_filename}]")
                except Exception as img_e:
                    print(f"Error extracting image: {img_e}")
            
            # Check if page has meaningful text content
            if not page_text.strip() or len(page_text.strip()) < 50:
                print(f"Page {page_num+1} appears to be a scanned image. Using OCR...")
                # Get the page as an image
                pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                img_data = pix.tobytes("png")
                
                # Convert bytes to image for OCR
                image = Image.open(io.BytesIO(img_data))
                
                # Save the page image for reference
                page_img_filename = f"{image_prefix}page{page_num+1}_full.png"
                page_img_path = os.path.join(images_dir, page_img_filename)
                image.save(page_img_path)
                
                # Use OCR to extract text
                ocr_text = pytesseract.image_to_string(image)
                text_content.append(ocr_text)
            else:
                text_content.append(page_text)
        
        return "\n".join(text_content), image_paths
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return "", []

def extract_text_from_image(image_path):
    """Extract text from image using pytesseract"""
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        print(f"Error extracting text from image: {e}")
        return ""

def improve_text_with_gemini(text):
    """Use Gemini AI to improve and structure the extracted text"""
    try:
        if not GOOGLE_API_KEY:
            return text, "API key not configured. Using raw extracted text."
          # Initialize Gemini model with the specific model name provided by the user
        model = genai.GenerativeModel('gemini-2.0-flash')
        # Create prompt for Gemini
        # Limiting text length to avoid exceeding model token limits
        text_for_prompt = text[:25000]
        prompt = f"""
        I'm providing text extracted from a document that contains multiple-choice questions from an exam. 
        Please help me structure and format it properly for a Word document.
        
        SPECIFIC FORMATTING REQUIREMENTS:
        1. Preserve the question numbering format (like "Câu 1 [ID]:")
        2. Format each multiple-choice question with the options (A, B, C, D) displayed as a list with proper spacing
        3. Properly format scientific formulas (e.g., convert O2 to O₂, CO2 to CO₂)
        4. Maintain the section headings (like "PHẦN I", "PHẦN II")
        5. Format any emphasized or italicized text (like gene names - "Gene Y", "Gene Z")
        6. Preserve any bullet points, numbered lists, and paragraph structure
        7. Make sure mathematical symbols are properly formatted
        
        VERY IMPORTANT: Return ONLY the corrected text content. DO NOT include any explanations, 
        instructions, or comments about what you've done. DO NOT include any headings like 
        "Key Improvements" or "How to Use in Word". DO NOT add any introduction or conclusion.
        JUST RETURN THE CLEAN, CORRECTED DOCUMENT TEXT.
        
        Here's the extracted text:
        
        {text_for_prompt}
        """
        
        # Generate response from Gemini
        response = model.generate_content(prompt)
        return response.text, "Text enhanced with AI"
    except Exception as e:
        print(f"Error using AI: {e}")
        return text, f"Error using AI: {str(e)}"

def create_word_document_with_images(text, output_path, image_paths):
    """Create a Word document from the improved text with better formatting"""
    try:
        doc = Document()
        
        # Set font for the entire document
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = docx.shared.Pt(11)
        
        # Prepare heading styles
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = 'Calibri'
        heading1_style.font.size = docx.shared.Pt(16)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = docx.shared.RGBColor(0, 92, 54)  # Green color
        
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = docx.shared.Pt(14)
        heading2_style.font.bold = True
        
        # Create a list style for multiple choice options
        list_style = doc.styles.add_style('Multiple Choice', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        list_style.font.name = 'Calibri'
        list_style.font.size = docx.shared.Pt(11)
        list_style.paragraph_format.left_indent = docx.shared.Pt(18)
        list_style.paragraph_format.first_line_indent = docx.shared.Pt(-18)
        
        # Format scientific formulas
        def format_scientific_formulas(text):
            # Replace common scientific formulas
            text = re.sub(r'O2\b', 'O₂', text)
            text = re.sub(r'CO2\b', 'CO₂', text)
            text = re.sub(r'H2O\b', 'H₂O', text)
            text = re.sub(r'H\+', 'H⁺', text)
            text = re.sub(r'([A-Za-z])(\d+)', r'\1₍\2₎', text)  # Convert normal digits to subscript
            return text
        
        # Process text by lines
        lines = text.split('\n')
        i = 0
        
        in_question = False
        current_question = None
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:  # Skip empty lines
                i += 1
                continue
            
            # Detect and format section headings (PHẦN I, PHẦN II, etc.)
            if re.match(r'^PHẦN\s+[IVX]+\.?', line, re.IGNORECASE):
                doc.add_heading(line, level=1)
                i += 1
                continue
            
            # Detect and format question numbering (Câu 1 [ID]:)
            question_match = re.match(r'^(Câu\s+\d+\s*\[\d+\]:?)(.*)', line, re.IGNORECASE)
            if question_match:
                question_num = question_match.group(1)
                question_text = question_match.group(2).strip()
                
                # Create a new paragraph for the question
                p = doc.add_paragraph()
                question_run = p.add_run(question_num + " ")
                question_run.bold = True
                
                # Add the question text with proper formula formatting
                question_text = format_scientific_formulas(question_text)
                p.add_run(question_text)
                
                in_question = True
                i += 1
                continue
            
            # Detect and format multiple choice options (A., B., C., D.)
            option_match = re.match(r'^([A-D])\.\s*(.*)', line)
            if option_match and in_question:
                option_letter = option_match.group(1)
                option_text = option_match.group(2).strip()
                
                # Format the option text
                option_text = format_scientific_formulas(option_text)
                
                # Check for italic text (like Gene Y, Gene Z)
                if re.search(r'Gene\s+[A-Z]', option_text):
                    p = doc.add_paragraph(style='Multiple Choice')
                    p.add_run(option_letter + ". ").bold = True
                    
                    # Split the text to apply italic only to the gene name
                    option_parts = re.split(r'(Gene\s+[A-Z])', option_text)
                    for part in option_parts:
                        if re.match(r'Gene\s+[A-Z]', part):
                            run = p.add_run(part)
                            run.italic = True
                        else:
                            p.add_run(part)
                else:
                    # Regular option without special formatting
                    p = doc.add_paragraph(style='Multiple Choice')
                    p.add_run(option_letter + ". ").bold = True
                    p.add_run(option_text)
                
                i += 1
                continue
            
            # Check for mathematical formulas - simple pattern detection
            if ('=' in line and (any(c in line for c in '+-*/^√∑∫πθ') or 
                              re.search(r'[a-zA-Z]_\d', line) or
                              re.search(r'\([^)]+\)', line))):
                p = doc.add_paragraph()
                formula_text = format_scientific_formulas(line)
                formula_run = p.add_run(formula_text)
                formula_run.italic = True
                formula_run.font.name = 'Cambria Math'  # Better font for math
                i += 1
                continue
            
            # Handle emphasized text with => or → symbol
            if "=>" in line or "→" in line:
                p = doc.add_paragraph()
                p.add_run(format_scientific_formulas(line)).bold = True
                i += 1
                continue
            
            # Normal paragraphs - collect consecutive non-empty lines
            paragraph_text = line
            j = i + 1
            while j < len(lines) and lines[j].strip() and not re.match(r'^[A-D]\.', lines[j]) and not re.match(r'^Câu\s+\d+', lines[j]):
                paragraph_text += ' ' + lines[j].strip()
                j += 1
            
            # Format and add the paragraph
            paragraph_text = format_scientific_formulas(paragraph_text)
            doc.add_paragraph(paragraph_text)
            
            # If we've moved beyond the current question
            if re.match(r'^Câu\s+\d+', paragraph_text):
                in_question = True
            
            i = j
        
        # Add images
        for image_path in image_paths:
            doc.add_picture(image_path)
        
        # Add some document properties
        core_properties = doc.core_properties
        core_properties.title = os.path.basename(output_path).split('.')[0]
        core_properties.author = "TW Converter with AI"
        core_properties.language = "vi-VN"  # Vietnamese
        
        # Save with some additional options for better compatibility
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return False

@app.route('/')
def index():
    """Render the homepage"""
    return render_template('index_minimal.html')  # Sử dụng template tối giản
    
@app.route('/original')
def original_page():
    """Render the original page"""
    return render_template('index.html')
    
@app.route('/test')
def test_page():
    """Render a simple test page to verify Flask is working properly"""
    return render_template('index_basic.html')

@app.route('/test-page')
def test_page_new():
    """Render a simple test page to verify Flask is working properly"""
    return render_template('test.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and conversion"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        # Generate unique filename
        original_filename = secure_filename(file.filename)
        filename_base = str(uuid.uuid4())
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename_base}.pdf")
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename_base}.docx")
        
        # Save the uploaded file
        file.save(pdf_path)
        
        # Extract text from PDF
        extracted_text, image_paths = extract_text_from_pdf(pdf_path)
        
        if not extracted_text:
            os.remove(pdf_path)  # Clean up
            return jsonify({'error': 'Failed to extract text from the PDF'}), 400
        
        # Improve text with AI
        improved_text, status = improve_text_with_gemini(extracted_text)
        
        # Create Word document
        if create_word_document_with_images(improved_text, docx_path, image_paths):
            # Create download URL
            download_url = url_for('download_file', filename=f"{filename_base}.docx", original_filename=original_filename.replace('.pdf', '.docx'))
            return jsonify({
                'success': True,
                'message': 'PDF successfully converted to Word',
                'status': status,
                'downloadUrl': download_url
            })
        else:
            os.remove(pdf_path)  # Clean up
            return jsonify({'error': 'Failed to create Word document'}), 500
    
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/upload_image', methods=['POST'])
def upload_image_file():
    """Handle image file upload and conversion"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if file and allowed_image_file(file.filename):
        # Generate unique filename
        original_filename = secure_filename(file.filename)
        filename_base = str(uuid.uuid4())
        image_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename_base}.{file.filename.rsplit('.', 1)[1].lower()}")
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename_base}.docx")
        output_filename = original_filename.rsplit('.', 1)[0] + '.docx'

        # Save the uploaded image
        file.save(image_path)

        # Extract text from image
        extracted_text = extract_text_from_image(image_path)

        if not extracted_text:
            os.remove(image_path)  # Clean up
            return jsonify({'error': 'Could not extract text from image'}), 400

        # Improve text using Gemini
        improved_text, ai_message = improve_text_with_gemini(extracted_text)

        # Create Word document
        success = create_word_document_with_images(improved_text, docx_path, [image_path])

        if not success:
            os.remove(image_path)  # Clean up
            return jsonify({'error': 'Failed to create Word document'}), 500

        # Return success with download URL
        return jsonify({
            'success': True,
            'message': ai_message,
            'downloadUrl': url_for('download_file', filename=f"{filename_base}.docx", original_filename=output_filename),
            'original_filename': output_filename
        })

    return jsonify({'error': 'Invalid file format. Only JPG, JPEG, and PNG files are allowed.'}), 400

@app.route('/download/<filename>')
def download_file(filename):
    """Handle file download"""
    original_filename = request.args.get('original_filename')
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, 
                             as_attachment=True, 
                             download_name=original_filename if original_filename else filename)

@app.route('/download-batch', methods=['GET'])
def download_batch():
    """Create a ZIP file with multiple converted documents and return it for download"""
    try:
        # Get list of files from query parameters
        files_param = request.args.get('files')
        if not files_param:
            return jsonify({'error': 'No files specified'}), 400
            
        # Add debug log
        print(f"Files parameter received: {files_param}")
            
        # Parse the JSON-encoded list of file URLs
        try:
            file_urls = json.loads(files_param)
            if not file_urls or not isinstance(file_urls, list):
                return jsonify({'error': 'Invalid files parameter'}), 400
                
            # Add debug log
            print(f"Parsed URLs: {file_urls}")
        except json.JSONDecodeError as e:
            print(f"JSON decode error: {e}")
            return jsonify({'error': f'Invalid JSON format in files parameter: {str(e)}'}), 400
        
        # Get file type (pdf or image)
        file_type = request.args.get('type', 'pdf')
        
        # Collect valid files
        valid_files = []
        
        for url in file_urls:
            # Skip if URL is None or not a string
            if not url or not isinstance(url, str):
                print(f"Skipping invalid URL: {url}")
                continue
                
            # Extract filename from URL
            try:
                # Process URL to extract filename
                if '/' in url:
                    url_parts = url.split('?')[0].split('/')
                    if len(url_parts) < 2:
                        print(f"URL doesn't have enough parts: {url}")
                        continue
                        
                    filename = url_parts[-1]
                else:
                    # Direct filename reference
                    filename = url.split('?')[0]
                
                # Print debug info
                print(f"Extracted filename: {filename}")
                
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                print(f"Looking for file at: {file_path}")
                
                # Check if file exists
                if os.path.exists(file_path):
                    print(f"File exists: {file_path}")
                    # Try to get original filename from query parameters
                    original_filename = None
                    if '?' in url and 'original_filename=' in url:
                        try:
                            query = url.split('?')[1]
                            for param in query.split('&'):
                                if param.startswith('original_filename='):
                                    original_filename = param.split('=')[1]
                                    # URL decode
                                    from urllib.parse import unquote
                                    original_filename = unquote(original_filename)
                                    print(f"Found original filename: {original_filename}")
                                    break
                        except Exception as parse_error:
                            print(f"Error parsing URL query: {parse_error}")
                    
                    # Add to valid files list
                    valid_files.append({
                        'path': file_path,
                        'original_name': original_filename if original_filename else filename
                    })
                else:
                    print(f"File not found: {file_path}")
            except Exception as url_error:
                print(f"Error processing URL {url}: {url_error}")
                continue
        
        # Check if we have any valid files
        if not valid_files:
            print("No valid files found")
            return jsonify({'error': 'No valid files found for download'}), 404
        
        print(f"Valid files: {valid_files}")
        
        # If only one file, return it directly
        if len(valid_files) == 1:
            file_info = valid_files[0]
            file_basename = os.path.basename(file_info['path'])
            print(f"Sending single file: {file_basename} as {file_info['original_name']}")
            return send_from_directory(
                app.config['UPLOAD_FOLDER'],
                file_basename,
                as_attachment=True,
                download_name=file_info['original_name']
            )
        
        # Create a unique zip filename
        zip_filename = f"tw_converter_{file_type}_{uuid.uuid4()}.zip"
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], zip_filename)
        print(f"Creating ZIP at: {zip_path}")
        
        # Create ZIP file
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file_info in valid_files:
                try:
                    # Debug message
                    print(f"Adding to ZIP: {file_info['path']} as {file_info['original_name']}")
                    
                    # Add file to ZIP
                    zipf.write(file_info['path'], arcname=file_info['original_name'])
                except Exception as zip_error:
                    print(f"Error adding file to ZIP: {zip_error}")
        
        # Check if the zip file was created successfully and has content
        if not os.path.exists(zip_path):
            print("ZIP file creation failed")
            return jsonify({'error': 'Failed to create ZIP file'}), 500
            
        zip_size = os.path.getsize(zip_path)
        print(f"ZIP file size: {zip_size} bytes")
        
        if zip_size < 100:  # Arbitrary small size check
            print("ZIP file is too small")
            return jsonify({'error': 'ZIP file seems to be empty or invalid'}), 500
            
        # Return the ZIP file
        print(f"Sending ZIP file: {zip_filename}")
        return send_from_directory(
            app.config['UPLOAD_FOLDER'], 
            zip_filename, 
            as_attachment=True, 
            download_name=f"tw_converter_{file_type}_files.zip"
        )
    except Exception as e:
        print(f"Error in download_batch: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/cleanup', methods=['POST'])
def cleanup_files():
    """Clean up temporary files"""
    data = request.get_json()
    if 'filename' in data:
        filename = data['filename']
        # Ensure the filename is safe and within our uploads directory
        if '..' not in filename and '/' not in filename:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
            
            # Remove both PDF and DOCX files
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                return jsonify({'success': True})
            except Exception as e:
                return jsonify({'error': str(e)}), 500
    
    return jsonify({'error': 'Invalid filename'}), 400

# Error handlers
@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 413

@app.errorhandler(500)
def internal_server_error(error):
    return jsonify({'error': 'Server error. Please try again later.'}), 500

@app.after_request
def add_header(response):
    """Add headers to handle caching appropriately for different types of content"""
    if 'text/html' in response.headers['Content-Type']:
        # Không cache cho HTML
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    elif 'javascript' in response.headers['Content-Type']:
        # Cache JavaScript files for 1 hour, but verify with server if changed
        response.headers['Cache-Control'] = 'public, max-age=3600, must-revalidate'
        response.headers['Vary'] = 'Accept-Encoding'
    elif 'image' in response.headers['Content-Type']:
        # Cache images for 1 day
        response.headers['Cache-Control'] = 'public, max-age=86400'
        
    # Add CORS headers for development
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    
    return response

if __name__ == '__main__':
    print("Starting Flask server...")
    print("Open your browser and navigate to http://127.0.0.1:5000/")
    try:
        # Sử dụng 127.0.0.1 thay vì 0.0.0.0 để tránh vấn đề với tường lửa
        app.run(debug=True, host='127.0.0.1', port=5000)
    except Exception as e:
        print(f"Error starting Flask server: {e}")
        print("Try running with default settings:")
        app.run(debug=True)
